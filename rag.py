"""
Minimal RAG pipeline you can read top to bottom.

The running scenario: Maya runs a bakery and has a binder of recipes and
allergy notes. The model has never seen that binder. RAG is how we let the
model answer questions about it without retraining anything.

Flow:  load -> chunk -> embed -> [save to disk] -> retrieve -> build prompt -> generate
                                       ^
                              next run: load from disk, skip embed

Two functions are left for YOU to write (marked "YOUR TURN"). They are the
two ideas that make RAG actually work: chunking and retrieval. Everything
else is done.
"""

from sentence_transformers import SentenceTransformer, CrossEncoder
import numpy as np
import json
import os
import torch
import pypdf
import docx
import openpyxl
import faiss

DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

EMBEDDER = SentenceTransformer("BAAI/bge-small-en-v1.5", device=DEVICE)
# BGE requires this prefix on queries only (not chunks) for retrieval tasks.
BGE_QUERY_PREFIX = "Represent this sentence for searching relevant passages: "

_RERANKER = None


def _get_reranker() -> CrossEncoder:
    global _RERANKER
    if _RERANKER is None:
        _RERANKER = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2", device=DEVICE)
    return _RERANKER


def _load_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _load_pdf(path: str) -> str:
    reader = pypdf.PdfReader(path)
    text_parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            text_parts.append(t)
    return "\n".join(text_parts)


def _load_docx(path: str) -> str:
    doc = docx.Document(path)
    text_parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(" | ".join(row_text))
    return "\n".join(text_parts)


def _load_excel(path: str) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    text_parts = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"Sheet: {sheet_name}")
        for row in sheet.iter_rows(values_only=True):
            row_vals = [str(val).strip() for val in row if val is not None]
            if row_vals:
                text_parts.append(" | ".join(row_vals))
    return "\n".join(text_parts)


# dispatch table: add a new extension here to support a new format
_LOADERS = {
    ".txt":  _load_txt,
    ".md":   _load_txt,
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".xlsx": _load_excel,
}


def load_documents(path: str) -> str:
    """Load any supported file type. Dispatches by extension."""
    ext = os.path.splitext(path)[1].lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(_LOADERS)}")
    return loader(path)


def load_folder(folder: str) -> list[tuple[str, str]]:
    """Read all supported files from folder. Returns [(filename, text), ...].
    Sorted so index order is stable across runs."""
    docs = []
    for name in sorted(os.listdir(folder)):
        ext = os.path.splitext(name)[1].lower()
        if ext in _LOADERS:
            full_path = os.path.join(folder, name)
            try:
                docs.append((name, load_documents(full_path)))
            except Exception as e:
                print(f"Error loading {name}: {e}")
    return docs


def chunk_text(text: str, size: int = 180, overlap: int = 20) -> list[str]:
    """Split text into word-based chunks. size=180 words stays under the
    all-MiniLM-L6-v2 limit of 256 tokens (~180 words). overlap=20 words
    keeps context alive at chunk boundaries."""
    words = text.split()
    step = size - overlap
    chunks = []
    start = 0
    while start < len(words):
        chunks.append(" ".join(words[start:start + size]))
        start += step
    return chunks


def embed(texts: list[str]) -> np.ndarray:
    """Turn a list of strings into a matrix of vectors, one row per string.
    normalize_embeddings=True makes every vector length 1, which is what lets
    us use a plain dot product as cosine similarity later."""
    return EMBEDDER.encode(texts, normalize_embeddings=True)


def build_faiss(vectors: np.ndarray) -> faiss.Index:
    """Wrap the vectors in a FAISS inner-product index.
    Vectors are normalized, so dot product == cosine similarity.
    FAISS needs contiguous float32, so we coerce.

    Index type is chosen ONCE by corpus size, so we never re-architect:
      - under ~10k vectors: IndexFlatIP, exact + instant. HNSW here would only
        lose recall for no speed gain.
      - at scale: IndexHNSWFlat, approximate nearest-neighbour (the friend's
        trick), trades a little recall for big speed on millions of vectors."""
    vectors = np.ascontiguousarray(vectors, dtype="float32")
    d = vectors.shape[1]
    if vectors.shape[0] < 10000:
        index = faiss.IndexFlatIP(d)
    else:
        index = faiss.IndexHNSWFlat(d, 32, faiss.METRIC_INNER_PRODUCT)
        index.hnsw.efSearch = 64
    index.add(vectors)
    return index


def retrieve(query: str, chunks: list[str], index: faiss.IndexFlatIP, k: int = 3) -> list[str]:
    """Return the `k` chunks whose vectors are most similar to the query's vector.
    FAISS does the nearest-neighbour search instead of a manual numpy scan."""
    q = EMBEDDER.encode([BGE_QUERY_PREFIX + query], normalize_embeddings=True)
    q = np.ascontiguousarray(q, dtype="float32")
    scores, top_idx = index.search(q, k)
    return [chunks[i] for i in top_idx[0]]


def rerank(query: str, candidates: list[str], k: int = 3) -> list[str]:
    """Re-score candidates with a cross-encoder and keep the best `k`.
    The bi-encoder (FAISS stage) is fast but approximate; the cross-encoder
    reads the query and chunk together, so it judges relevance more precisely."""
    if not candidates:
        return []
    pairs = [(query, c) for c in candidates]
    scores = _get_reranker().predict(pairs)
    order = np.argsort(scores)[::-1][:k]
    return [candidates[i] for i in order]


def search(query: str, chunks: list[str], index: faiss.IndexFlatIP,
           k: int = 3, fetch_k: int = 10) -> list[str]:
    """Two-stage retrieval: FAISS casts a wide net (fetch_k), reranker sharpens it (k).
    This is the standard production pattern: cheap recall, then precise reorder."""
    candidates = retrieve(query, chunks, index, k=fetch_k)
    return rerank(query, candidates, k=k)


def save_index(chunks: list[str], vectors: np.ndarray, folder: str = "data/index") -> None:
    """Save chunks + vectors to disk so next run skips embedding entirely.
    Two files: chunks.json (the text) and vectors.npy (the numbers)."""
    os.makedirs(folder, exist_ok=True)
    with open(f"{folder}/chunks.json", "w") as f:
        json.dump(chunks, f)
    np.save(f"{folder}/vectors.npy", vectors)


def load_index(folder: str = "data/index") -> tuple[list[str], np.ndarray] | None:
    """Load from disk if the index exists. Returns None if not built yet.
    Caller decides: None means embed and save; tuple means use directly."""
    chunks_path = f"{folder}/chunks.json"
    vectors_path = f"{folder}/vectors.npy"
    if not (os.path.exists(chunks_path) and os.path.exists(vectors_path)):
        return None
    with open(chunks_path) as f:
        chunks = json.load(f)
    vectors = np.load(vectors_path)
    return chunks, vectors


def build_prompt(query: str, context_chunks: list[str]) -> str:
    """Stuff the retrieved chunks into the prompt as grounding context.
    The 'use ONLY the context' instruction is what fights hallucination."""
    context = "\n\n".join(context_chunks)
    return (
        "Answer the question using ONLY the context below. "
        "If the answer is not in the context, say you do not know.\n\n"
        f"Context:\n{context}\n\n"
        f"Question: {query}\n"
        "Answer:"
    )


def generate(prompt: str) -> str:
    """Send the grounded prompt to a local Ollama model. No API key.
    Needs Ollama running and the model pulled: `ollama pull llama3.2`."""
    import ollama

    resp = ollama.chat(
        model="qwen3.5:4b",  # change to any model you have pulled
        messages=[{"role": "user", "content": prompt}],
    )
    return resp["message"]["content"]


def build_index(path: str) -> tuple[list[str], faiss.IndexFlatIP]:
    """Load or build the index. Returns chunks + a ready FAISS index.
    Vectors persist to disk; the FAISS index is rebuilt in memory each run
    (instant for IndexFlatIP, so no separate index file needed)."""
    cached = load_index()
    if cached:
        chunks, vectors = cached
        print("(loaded index from disk, skipped embedding)")
    else:
        if os.path.isdir(path):
            docs = load_folder(path)
        else:
            docs = [(os.path.basename(path), load_documents(path))]
        chunks = []
        for filename, text in docs:
            for chunk in chunk_text(text):
                chunks.append(f"[source: {filename}]\n{chunk}")
        vectors = embed(chunks)
        save_index(chunks, vectors)
        print(f"(built and saved index from {len(docs)} file(s))")
    return chunks, build_faiss(vectors)


def answer_question(path: str, query: str) -> str:
    """Single-shot answer. Builds index if needed, answers once."""
    chunks, index = build_index(path)
    top = search(query, chunks, index, k=3)
    return generate(build_prompt(query, top))


def repl(path: str) -> None:
    """Interactive loop. Index loads once; ask as many questions as you want.
    Type 'quit' or press Ctrl-C to exit."""
    chunks, index = build_index(path)
    print("Index ready. Ask anything (type 'quit' to exit).\n")
    while True:
        try:
            query = input("You: ").strip()
        except (KeyboardInterrupt, EOFError):
            print("\nBye.")
            break
        if not query or query.lower() == "quit":
            break
        top = search(query, chunks, index, k=3)
        print(f"\nRAG: {generate(build_prompt(query, top))}\n")


def _selfcheck() -> None:
    """Runs the pipeline up to retrieval. No API key needed. Run: python rag.py check"""
    text = load_documents("data/maya_binder.txt")
    chunks = chunk_text(text, size=30, overlap=5)
    assert len(chunks) > 1, "chunking should produce multiple pieces"
    assert chunks[1].split()[0] in text, "chunks must be words from the text"
    index = build_faiss(embed(chunks))
    top = retrieve("which cake is safe for a nut allergy", chunks, index, k=1)
    assert "nut" in top[0].lower() or "almond" in top[0].lower(), "retrieval missed the nut chunk"
    print("OK. Top chunk for the nut-allergy query:\n")
    print(top[0])


if __name__ == "__main__":
    import sys
    path = "data/maya_binder.txt"
    if len(sys.argv) > 1 and sys.argv[1] == "check":
        _selfcheck()
    elif len(sys.argv) > 1 and sys.argv[1] == "chat":
        repl(sys.argv[2] if len(sys.argv) > 2 else path)
    else:
        print(answer_question(path, "Which cake is safe for a nut allergy?"))
