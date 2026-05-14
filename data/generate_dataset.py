import os
import docx
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def create_txt(directory):
    path = os.path.join(directory, "bakery_intro.txt")
    content = """Maya's Sweet Haven - Welcome Guide and Operations
Established in 2024, Maya's Sweet Haven is a craft bakery specializing in organic, traditional, and allergy-conscious baked goods.
Location: 123 Baker Street, Sweetwater.
Hours of Operation:
- Monday to Friday: 7:00 AM - 6:00 PM
- Saturday: 8:00 AM - 4:00 PM
- Sunday: Closed

Mission Statement:
Our mission is to bring joy through baking. We use only locally sourced, organic ingredients, and we strive to provide safe, delicious options for customers with dietary restrictions, including gluten-free, nut-free, and vegan products. All of our bakers are certified in allergen cross-contamination prevention."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created {path}")

def create_docx(directory):
    path = os.path.join(directory, "recipes_and_allergies.docx")
    doc = docx.Document()
    doc.add_heading("Maya's Sweet Haven Recipe Binder - Secrets & Safety", level=1)
    
    doc.add_heading("Chocolate Lava Cake", level=2)
    doc.add_paragraph("Ingredients: 200g dark chocolate, 100g butter, 100g sugar, 4 eggs, 50g flour.")
    doc.add_paragraph("Instructions: Melt chocolate and butter. Whisk eggs and sugar. Fold in flour and melted mixture. Bake at 200C for 10 minutes.")
    doc.add_paragraph("Allergen Warning: Contains eggs, dairy, and gluten (wheat flour).")
    
    doc.add_heading("Gluten-Free Almond Macarons", level=2)
    doc.add_paragraph("Ingredients: 150g almond flour, 150g powdered sugar, 3 egg whites, 100g granulated sugar.")
    doc.add_paragraph("Instructions: Whisk egg whites to stiff peaks, folding in almond flour and sugar. Pipe onto tray, let sit for 30 minutes to form skin. Bake at 150C for 15 minutes.")
    doc.add_paragraph("Allergen Warning: Contains almonds (tree nuts) and eggs. Strictly gluten-free.")
    
    doc.add_heading("Vegan Blueberry Muffins", level=2)
    doc.add_paragraph("Ingredients: 250g flour, 2 tsp baking powder, 100g sugar, 200ml almond milk, 75ml vegetable oil, 150g fresh blueberries.")
    doc.add_paragraph("Instructions: Mix dry ingredients. Whisk wet ingredients. Combine and gently fold in blueberries. Bake at 180C for 20 minutes.")
    doc.add_paragraph("Allergen Warning: Contains almonds (in the almond milk). Gluten-present.")
    
    doc.save(path)
    print(f"Created {path}")

def create_xlsx(directory):
    path = os.path.join(directory, "inventory_pricing.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Product Inventory"
    
    headers = ["Product ID", "Product Name", "Category", "Price", "Primary Allergen", "Status"]
    ws.append(headers)
    
    rows = [
        ["PROD-001", "Chocolate Lava Cake", "Cake", "$6.50", "Gluten, Dairy, Eggs", "Active"],
        ["PROD-002", "Almond Macarons", "Cookie", "$4.00", "Tree Nuts, Eggs", "Active"],
        ["PROD-003", "Vegan Blueberry Muffin", "Muffin", "$3.70", "Tree Nuts", "Active"],
        ["PROD-004", "Classic Croissant", "Pastry", "$3.50", "Gluten, Dairy", "Active"],
        ["PROD-005", "Gluten-Free Bread", "Bread", "$5.50", "None", "Active"]
    ]
    for r in rows:
        ws.append(r)
        
    wb.save(path)
    print(f"Created {path}")

def create_pdf(directory):
    path = os.path.join(directory, "safety_compliance.pdf")
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    
    # Page 1
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, "Maya's Sweet Haven - Kitchen Safety & Compliance Manual")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 100, "1. Allergen Separation Protocol")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 120, "To prevent cross-contamination, all gluten-free and nut-free items must be prepared")
    c.drawString(50, height - 135, "on dedicated color-coded prep tables. Purple boards and utensils are for allergen-free prep.")
    c.drawString(50, height - 150, "Bakers must wash hands and change aprons before switching to allergen-free recipes.")
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 200, "2. Ingredient Storage Rules")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 220, "All tree nut ingredients (almond flour, walnuts, pecans) must be stored in airtight,")
    c.drawString(50, height - 235, "labeled containers on the bottom shelf of the dry storage rack to avoid spillages onto")
    c.drawString(50, height - 250, "allergen-safe ingredients below.")
    
    c.showPage()
    
    # Page 2
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 50, "3. Emergency Allergy Response")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 70, "In the event of a customer reporting an allergic reaction:")
    c.drawString(50, height - 85, "- Locate the First Aid Kit in the main kitchen near the emergency exit.")
    c.drawString(50, height - 100, "- Call emergency services (911) immediately if the customer shows signs of anaphylaxis.")
    c.drawString(50, height - 115, "- Administer the store-provided epinephrine auto-injector if authorized and necessary.")
    c.drawString(50, height - 130, "- Report the incident immediately to Maya or the manager on duty.")
    
    c.save()
    print(f"Created {path}")

def main():
    directory = "data/dataset"
    os.makedirs(directory, exist_ok=True)
    create_txt(directory)
    create_docx(directory)
    create_xlsx(directory)
    create_pdf(directory)

if __name__ == "__main__":
    main()
